import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc_path = 'Ai text Detector Report_Updated.docx'
    doc = docx.Document(doc_path)
    
    idx_seq = None
    for i, p in enumerate(doc.paragraphs):
        if 'Figure 3.6: Sequence Diagram for AI Text Detector' in p.text:
            idx_seq = i
            break

    if idx_seq is not None:
        p_img = doc.paragraphs[idx_seq - 1]
        p_img.text = ""
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_img.add_run()
        r.add_picture('docs/figure_3_6_sequence_diagram.png', width=Inches(5.8))

        p_exp = doc.paragraphs[idx_seq + 1]
        p_exp.text = (
            "Figure 3.6 demonstrates the streamlined sequence diagram for the AI Text Detector system, "
            "designed in a clean Draw.io layout across four principal lifelines: Client, Server, Scanner Engine, and Database. "
            "The operational flow is executed in nine sequential steps:\n\n"
            "1. Client -> Server: The user submits a text passage or drops a document (.docx, .pdf, .txt) via the dashboard GUI (Step 1: Scan request).\n"
            "2. Server -> Scanner Engine: The server controller validates word count threshold (minimum 10 words) and triggers the detection engine (Step 2: Initiate scan).\n"
            "3. Scanner Engine <-> Database: The engine queries pre-compiled feature profiles, training centroids, and n-gram models from SQLite/IndexedDB storage (Step 3: Payloads / models).\n"
            "4. Scanner Engine (Loop): Inside the 'Scan all AI evaluation engines' loop box, parallel feature extraction and classification are performed across Perplexity, K-NN (K=5), Cosine Similarity, POS syntax, Trigram, and Grammar Calibration (Step 4).\n"
            "5. Scanner Engine -> Server: The classification ensemble computes the weighted score and passes back component metrics and overall probability (Step 5: Return).\n"
            "6. Scanner Engine -> Database: Scan metadata, extracted features, and final AI probability are persisted to local history (Step 6: Store results).\n"
            "7. Server -> Scanner Engine: The controller requests formatted SVG chart rendering payloads (Step 7: Request report).\n"
            "8. Scanner Engine -> Server: Visual rendering data for the radial speedometer gauge, scatter plot, radar web, and sentence heatmap are generated (Step 8: Send report).\n"
            "9. Server -> Client: The server delivers the complete interactive dashboard view and risk verdict to the user interface (Step 9: Return report)."
        )
        
        doc.save(doc_path)
        print("Successfully updated Sequence Diagram image and explanation in report!")

if __name__ == '__main__':
    main()
