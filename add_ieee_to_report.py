import docx
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml

def create_ieee_table(doc, algo_num, title, inputs, outputs, lines):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    cell = tbl.cell(0, 0)
    
    # Border & Shading XML
    tcPr = cell._element.get_or_add_tcPr()
    borders_xml = parse_xml('''
        <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:val="single" w:sz="14" w:space="0" w:color="003399"/>
            <w:left w:val="none"/>
            <w:bottom w:val="single" w:sz="14" w:space="0" w:color="003399"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders_xml)
    
    shd_xml = parse_xml('''<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="F8F9FA"/>''')
    tcPr.append(shd_xml)
    
    # Algorithm Header Title
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_before = Pt(4)
    p0.paragraph_format.space_after = Pt(4)
    r_title = p0.add_run(f"Algorithm {algo_num}: {title}")
    r_title.bold = True
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(10.5)
    r_title.font.color.rgb = RGBColor(0, 51, 153)
    
    # Divider line 1
    p_div1 = cell.add_paragraph()
    p_div1.paragraph_format.space_before = Pt(0)
    p_div1.paragraph_format.space_after = Pt(4)
    p_div1_xml = parse_xml('''
        <w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>
        </w:pBdr>
    ''')
    p_div1._element.get_or_add_pPr().append(p_div1_xml)
    
    # Input / Output
    p_io = cell.add_paragraph()
    p_io.paragraph_format.space_before = Pt(2)
    p_io.paragraph_format.space_after = Pt(4)
    p_io.paragraph_format.line_spacing = 1.15
    
    r_in = p_io.add_run("Input: ")
    r_in.bold = True
    r_in.font.name = 'Calibri'
    r_in.font.size = Pt(9.5)
    r_in_txt = p_io.add_run(inputs + "\n")
    r_in_txt.font.name = 'Calibri'
    r_in_txt.font.size = Pt(9.5)
    
    r_out = p_io.add_run("Output: ")
    r_out.bold = True
    r_out.font.name = 'Calibri'
    r_out.font.size = Pt(9.5)
    r_out_txt = p_io.add_run(outputs)
    r_out_txt.font.name = 'Calibri'
    r_out_txt.font.size = Pt(9.5)
    
    # Divider line 2
    p_div2 = cell.add_paragraph()
    p_div2.paragraph_format.space_before = Pt(0)
    p_div2.paragraph_format.space_after = Pt(4)
    p_div2_xml = parse_xml('''
        <w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>
        </w:pBdr>
    ''')
    p_div2._element.get_or_add_pPr().append(p_div2_xml)
    
    # Lines of Pseudocode
    for line_num, indent_level, line_tokens in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.15
        
        r_num = p.add_run(f"{line_num:2d}:  ")
        r_num.font.name = 'Consolas'
        r_num.font.size = Pt(9)
        r_num.font.color.rgb = RGBColor(120, 120, 120)
        
        if indent_level > 0:
            r_ind = p.add_run("    " * indent_level)
            r_ind.font.name = 'Consolas'
            r_ind.font.size = Pt(9)
        
        for text, is_kw, is_comment, is_math in line_tokens:
            r = p.add_run(text)
            r.font.name = 'Consolas'
            r.font.size = Pt(9)
            if is_kw:
                r.bold = True
                r.font.color.rgb = RGBColor(0, 51, 153) # Navy blue
            elif is_comment:
                r.italic = True
                r.font.color.rgb = RGBColor(110, 110, 110) # Muted grey
            elif is_math:
                r.font.color.rgb = RGBColor(153, 0, 77) # Magenta/wine
            else:
                r.font.color.rgb = RGBColor(30, 30, 30)
                
    return tbl


def main():
    doc = docx.Document('Ai text Detector Report_Updated.docx')
    print("Loaded report document. Total paragraphs:", len(doc.paragraphs))
    
    # -------------------------------------------------------------
    # ALGORITHM 1
    # -------------------------------------------------------------
    lines_algo1 = [
        (1, 0, [("procedure ", True, False, False), ("STATISTICALNLPANALYSIS", False, False, True), ("(T)", False, False, False)]),
        (2, 1, [("sentences ", False, False, True), ("<- ", False, False, False), ("SplitSentences", False, False, True), ("(T)", False, False, False), ("  // Regex split via lookbehind", False, True, False)]),
        (3, 1, [("words ", False, False, True), ("<- ", False, False, False), ("TokenizeWords", False, False, True), ("(Lower(T))", False, False, False)]),
        (4, 1, [("if ", True, False, False), ("|words| < 5 ", False, False, False), ("then return ", True, False, False), ("50.0  ", False, False, True), ("// Short document fallback", False, True, False)]),
        (5, 1, [("// Phase 1: Sentence Length Variance (Burstiness)", False, True, False)]),
        (6, 1, [("mu ", False, False, True), ("<- (1 / |sentences|) * sum(L_i)", False, False, False)]),
        (7, 1, [("sigma ", False, False, True), ("<- sqrt((1 / |sentences|) * sum((L_i - mu)^2))", False, False, False)]),
        (8, 1, [("// Phase 2: Lexical Diversity (Type-Token Ratio)", False, True, False)]),
        (9, 1, [("TTR ", False, False, True), ("<- |UniqueWords| / |words|", False, False, False)]),
        (10, 1, [("// Phase 3: Shannon Entropy & Predictability Penalty", False, True, False)]),
        (11, 1, [("H(X) ", False, False, True), ("<- -sum(P(w) * log2(P(w)))", False, False, False)]),
        (12, 1, [("penalty ", False, False, True), ("<- (bigramRatio * 50) + (aiWordRatio * 250)", False, False, False)]),
        (13, 1, [("// Phase 4: Component Normalization & Weighted Combination", False, True, False)]),
        (14, 1, [("burstFactor ", False, False, True), ("<- clamp(100 - (sigma * 9), 0, 100)", False, False, False)]),
        (15, 1, [("TTR_exp ", False, False, True), ("<- 0.86 - (|words| * 0.0003)", False, False, False)]),
        (16, 1, [("ttrFactor ", False, False, True), ("<- clamp(50 + (TTR_exp - TTR) * 250, 0, 100)", False, False, False)]),
        (17, 1, [("predFactor ", False, False, True), ("<- clamp((penalty / 18) * 100, 0, 100)", False, False, False)]),
        (18, 1, [("S_Perplexity ", False, False, True), ("<- (burstFactor * 0.30) + (ttrFactor * 0.20) + (predFactor * 0.50)", False, False, False)]),
        (19, 1, [("return ", True, False, False), ("S_Perplexity", False, False, True)]),
        (20, 0, [("end procedure", True, False, False)])
    ]
    
    # -------------------------------------------------------------
    # ALGORITHM 2
    # -------------------------------------------------------------
    lines_algo2 = [
        (1, 0, [("procedure ", True, False, False), ("KNNCLASSIFY", False, False, True), ("(T, D)", False, False, False)]),
        (2, 1, [("// Phase 1: 5D Feature Vector Extraction & Normalization", False, True, False)]),
        (3, 1, [("f0 ", False, False, True), ("<- clamp(0.5 + (TTR_exp - TTR) * 5.0, 0, 1)  // Lexical Diversity", False, False, False)]),
        (4, 1, [("f1 ", False, False, True), ("<- clamp(1.0 - (burstiness / 30.0), 0, 1)  // Sentence Burstiness", False, False, False)]),
        (5, 1, [("f2 ", False, False, True), ("<- clamp((aiWordCount / |words|) / 0.012, 0, 1)  // AI Word Density", False, False, False)]),
        (6, 1, [("f3 ", False, False, True), ("<- clamp(1 - |13 - CLI| / 12, 0, 1)  // Readability Index", False, False, False)]),
        (7, 1, [("f4 ", False, False, True), ("<- clamp(1 - |11 - punctPer100| / 12, 0, 1)  // Punctuation Density", False, False, False)]),
        (8, 1, [("v ", False, False, True), ("<- [f0, f1, f2, f3, f4]", False, False, False)]),
        (9, 1, [("// Phase 2: Weighted Euclidean Distance Computation", False, True, False)]),
        (10, 1, [("w ", False, False, True), ("<- [2.5, 2.5, 1.5, 0.5, 0.5]", False, False, False)]),
        (11, 1, [("for each ", True, False, False), ("(x_j, y_j) in D ", False, False, True), ("do", True, False, False)]),
        (12, 2, [("d_j ", False, False, True), ("<- sqrt(sum(w_i * (v_i - x_j,i)^2))", False, False, False)]),
        (13, 1, [("end for", True, False, False)]),
        (14, 1, [("// Phase 3: Distance Sorting & Voting Among K=5 Neighbors", False, True, False)]),
        (15, 1, [("N_K ", False, False, True), ("<- SelectTopKSmallestDistances(D, K=5)", False, False, False)]),
        (16, 1, [("aiVotes ", False, False, True), ("<- |{ n in N_K : label(n) == AI }|", False, False, False)]),
        (17, 1, [("S_KNN ", False, False, True), ("<- (aiVotes / 5) * 100", False, False, False)]),
        (18, 1, [("return ", True, False, False), ("S_KNN, N_K", False, False, True)]),
        (19, 0, [("end procedure", True, False, False)])
    ]
    
    # -------------------------------------------------------------
    # ALGORITHM 3
    # -------------------------------------------------------------
    lines_algo3 = [
        (1, 0, [("procedure ", True, False, False), ("COSINESIMILARITYANALYSIS", False, False, True), ("(T)", False, False, False)]),
        (2, 1, [("// Phase 1: Stylometric Marker Feature Extraction", False, True, False)]),
        (3, 1, [("v0 ", False, False, True), ("<- min(1.0, firstPersonRatio / 0.04)  // 1st Person Pronouns", False, False, False)]),
        (4, 1, [("v1 ", False, False, True), ("<- min(1.0, transitionRatio / 0.04)  // Complex Conjunctions", False, False, False)]),
        (5, 1, [("v2 ", False, False, True), ("<- min(1.0, intensifierRatio / 0.045) // Intensifiers", False, False, False)]),
        (6, 1, [("v3 ", False, False, True), ("<- min(1.0, passiveVoiceRatio / 0.015) // Passive Voice", False, False, False)]),
        (7, 1, [("v4 ", False, False, True), ("<- clamp(0.3 + (hapax_exp - hapaxRatio) * 2, 0, 1) // Hapax", False, False, False)]),
        (8, 1, [("v ", False, False, True), ("<- [v0, v1, v2, v3, v4]", False, False, False)]),
        (9, 1, [("// Phase 2: Vector Space Dot-Product Cosine Evaluation", False, True, False)]),
        (10, 1, [("sim_AI ", False, False, True), ("<- (v . P_AI) / (||v|| * ||P_AI||)", False, False, False)]),
        (11, 1, [("sim_Human ", False, False, True), ("<- (v . P_Human) / (||v|| * ||P_Human||)", False, False, False)]),
        (12, 1, [("// Phase 3: Linear Difference Score Mapping", False, True, False)]),
        (13, 1, [("S_Cosine ", False, False, True), ("<- clamp(((sim_AI - sim_Human) + 1) / 2 * 100, 0, 100)", False, False, False)]),
        (14, 1, [("return ", True, False, False), ("S_Cosine", False, False, True)]),
        (15, 0, [("end procedure", True, False, False)])
    ]
    
    # -------------------------------------------------------------
    # ALGORITHM 4
    # -------------------------------------------------------------
    lines_algo4 = [
        (1, 0, [("procedure ", True, False, False), ("LOGISTICREGRESSIONPREDICT", False, False, True), ("(v)", False, False, False)]),
        (2, 1, [("// Phase 1: Learned Model Parameter Initialization", False, True, False)]),
        (3, 1, [("beta0 ", False, False, True), ("<- -10.87000  // Intercept coefficient", False, False, False)]),
        (4, 1, [("w ", False, False, True), ("<- [6.79418, -0.17427, 0.58039, 5.29141, 2.93657] // Learned Weights", False, False, False)]),
        (5, 1, [("// Phase 2: Linear Logit Combination", False, True, False)]),
        (6, 1, [("z ", False, False, True), ("<- beta0 + sum(w_i * v_i)", False, False, False)]),
        (7, 1, [("// Phase 3: Sigmoid Activation Transformation", False, True, False)]),
        (8, 1, [("P_AI ", False, False, True), ("<- 1 / (1 + exp(-z))", False, False, False)]),
        (9, 1, [("S_LR ", False, False, True), ("<- P_AI * 100", False, False, False)]),
        (10, 1, [("return ", True, False, False), ("S_LR", False, False, True)]),
        (11, 0, [("end procedure", True, False, False)])
    ]

    # -------------------------------------------------------------
    # ALGORITHM 5
    # -------------------------------------------------------------
    lines_algo5 = [
        (1, 0, [("procedure ", True, False, False), ("TRIGRAMANALYSIS", False, False, True), ("(T)", False, False, False)]),
        (2, 1, [("if ", True, False, False), ("|T| < 10 ", False, False, False), ("then return ", True, False, False), ("50.0  ", False, False, True), ("// Short text fallback", False, True, False)]),
        (3, 1, [("T_prime ", False, False, True), ("<- CollapseWhitespace(Lower(T))", False, False, False)]),
        (4, 1, [("// Phase 1: Sliding Window Character Trigram Counting", False, True, False)]),
        (5, 1, [("for ", True, False, False), ("i <- 0 ", False, False, True), ("to ", True, False, False), ("|T_prime| - 3 ", False, False, True), ("do", True, False, False)]),
        (6, 2, [("tg ", False, False, True), ("<- T_prime[i : i+3]", False, False, False)]),
        (7, 2, [("if ", True, False, False), ("tg in TargetTrigrams ", False, False, True), ("then ", True, False, False), ("counts[tg] <- counts[tg] + 1", False, False, False)]),
        (8, 1, [("end for", True, False, False)]),
        (9, 1, [("// Phase 2: Frequency Scaling (10,000x Factor)", False, True, False)]),
        (10, 1, [("for each ", True, False, False), ("tg_k in TargetTrigrams ", False, False, True), ("do", True, False, False)]),
        (11, 2, [("V_k ", False, False, True), ("<- (counts[tg_k] / totalTrigrams) * 10000", False, False, False)]),
        (12, 1, [("end for", True, False, False)]),
        (13, 1, [("// Phase 3: Profile Cosine Similarity Comparison", False, True, False)]),
        (14, 1, [("sim_AI ", False, False, True), ("<- (V . AI_ref) / (||V|| * ||AI_ref||)", False, False, False)]),
        (15, 1, [("sim_Human ", False, False, True), ("<- (V . H_ref) / (||V|| * ||H_ref||)", False, False, False)]),
        (16, 1, [("S_Trigram ", False, False, True), ("<- clamp(((sim_AI - sim_Human) + 1) / 2 * 100, 0, 100)", False, False, False)]),
        (17, 1, [("return ", True, False, False), ("S_Trigram", False, False, True)]),
        (18, 0, [("end procedure", True, False, False)])
    ]

    # -------------------------------------------------------------
    # ALGORITHM 6
    # -------------------------------------------------------------
    lines_algo6 = [
        (1, 0, [("procedure ", True, False, False), ("POSRATIONANALYSIS", False, False, True), ("(T)", False, False, False)]),
        (2, 1, [("words ", False, False, True), ("<- TokenizeWords(Lower(T))", False, False, False)]),
        (3, 1, [("Categories ", False, False, True), ("<- {Determiners, Prepositions, Pronouns, Conjunctions, AuxiliaryVerbs}", False, False, False)]),
        (4, 1, [("// Phase 1: Lexicon Matching & Syntactic Tallying", False, True, False)]),
        (5, 1, [("for each ", True, False, False), ("w in words ", False, False, True), ("do", True, False, False)]),
        (6, 2, [("for each ", True, False, False), ("cat in Categories ", False, False, True), ("do", True, False, False)]),
        (7, 3, [("if ", True, False, False), ("w in Lexicon(cat) ", False, False, True), ("then ", True, False, False), ("counts[cat] <- counts[cat] + 1", False, False, False)]),
        (8, 2, [("end for", True, False, False)]),
        (9, 1, [("end for", True, False, False)]),
        (10, 1, [("// Phase 2: Relative Frequency Vector Construction", False, True, False)]),
        (11, 1, [("for each ", True, False, False), ("cat in Categories ", False, False, True), ("do", True, False, False)]),
        (12, 2, [("v[cat] ", False, False, True), ("<- counts[cat] / |words|", False, False, False)]),
        (13, 1, [("end for", True, False, False)]),
        (14, 1, [("// Phase 3: Centroid Cosine Similarity Comparison", False, True, False)]),
        (15, 1, [("sim_AI ", False, False, True), ("<- (v . C_AI) / (||v|| * ||C_AI||)", False, False, False)]),
        (16, 1, [("sim_Human ", False, False, True), ("<- (v . C_Human) / (||v|| * ||C_Human||)", False, False, False)]),
        (17, 1, [("S_POS ", False, False, True), ("<- clamp(((sim_AI - sim_Human) + 1) / 2 * 100, 0, 100)", False, False, False)]),
        (18, 1, [("return ", True, False, False), ("S_POS", False, False, True)]),
        (19, 0, [("end procedure", True, False, False)])
    ]

    # -------------------------------------------------------------
    # ALGORITHM 7
    # -------------------------------------------------------------
    lines_algo7 = [
        (1, 0, [("procedure ", True, False, False), ("GRAMMARSTYLEANALYSIS", False, False, True), ("(T)", False, False, False)]),
        (2, 1, [("totalPenalty ", False, False, True), ("<- 0; ", False, False, False), ("I ", False, False, True), ("<- empty_set", False, False, False)]),
        (3, 1, [("// Phase 1: Rule Pattern Regex Evaluation", False, True, False)]),
        (4, 1, [("for each ", True, False, False), ("(regex_i, weight_i, name_i) in R ", False, False, True), ("do", True, False, False)]),
        (5, 2, [("matches_i ", False, False, True), ("<- FindAllMatches(regex_i, T)", False, False, False)]),
        (6, 2, [("totalPenalty ", False, False, True), ("<- totalPenalty + (|matches_i| * weight_i)", False, False, False)]),
        (7, 2, [("if ", True, False, False), ("|matches_i| > 0 ", False, False, False), ("then ", True, False, False), ("I <- I union {(name_i, matches_i)}", False, False, False)]),
        (8, 1, [("end for", True, False, False)]),
        (9, 1, [("// Phase 2: Penalty Normalization per 100 Words", False, True, False)]),
        (10, 1, [("normPenalty ", False, False, True), ("<- (totalPenalty / max(10, |words|)) * 100", False, False, False)]),
        (11, 1, [("// Phase 3: Perfection Score Calculation", False, True, False)]),
        (12, 1, [("G ", False, False, True), ("<- clamp(100 - (normPenalty * 15), 0, 100)", False, False, False)]),
        (13, 1, [("return ", True, False, False), ("G, I", False, False, True)]),
        (14, 0, [("end procedure", True, False, False)])
    ]

    # -------------------------------------------------------------
    # ALGORITHM 8
    # -------------------------------------------------------------
    lines_algo8 = [
        (1, 0, [("procedure ", True, False, False), ("ENSEMBLECALIBRATION", False, False, True), ("(S_LR, S_Trigram, S_KNN, S_POS, S_Cosine, S_Perplexity, G)", False, False, False)]),
        (2, 1, [("// Phase 1: Weighted Linear Sum Compilation", False, True, False)]),
        (3, 1, [("P_ensemble ", False, False, True), ("<- (S_LR * 0.25) + (S_Trigram * 0.25) + (S_KNN * 0.20)", False, False, False)]),
        (4, 1, [("             + (S_POS * 0.15) + (S_Cosine * 0.10) + (S_Perplexity * 0.05)", False, False, False)]),
        (5, 1, [("// Phase 2: Grammar Discount & Flawless Boost Calibration", False, True, False)]),
        (6, 1, [("grammarFactor ", False, False, True), ("<- G / 100.0", False, False, False)]),
        (7, 1, [("if ", True, False, False), ("G == 100.0 ", False, False, False), ("then ", True, False, False), ("perfectBoost <- 15.0", False, False, False)]),
        (8, 1, [("else ", True, False, False), ("perfectBoost <- 0.0", False, False, False)]),
        (9, 1, [("end if", True, False, False)]),
        (10, 1, [("// Phase 3: Final Bounded Calibration Resolution", False, True, False)]),
        (11, 1, [("P_final ", False, False, True), ("<- min(100.0, max(0.0, (P_ensemble * grammarFactor) + perfectBoost))", False, False, False)]),
        (12, 1, [("return ", True, False, False), ("P_final", False, False, True)]),
        (13, 0, [("end procedure", True, False, False)])
    ]

    # List of specifications
    specs = [
        ("3.3.1. Statistical NLP Algorithm", 1, "Statistical NLP Analysis (Perplexity & Burstiness)", "Document text string T", "AI Probability Score S_Perplexity in [0, 100], Per-sentence AI heatmap", lines_algo1),
        ("3.3.2. KNN Classifier", 2, "K-Nearest Neighbors Classification (K=5)", "Document text T, Training dataset D = {(x_j, y_j)} of 43 labeled 5D vectors", "AI Probability Score S_KNN in [0, 100], Top K=5 nearest neighbor metadata", lines_algo2),
        ("3.3.3. Cosine Similarity (Vector Space Model)", 3, "Stylometric Cosine Similarity Analysis", "Document text T, Prototype centroid vectors P_AI, P_Human in R^5", "AI Probability Score S_Cosine in [0, 100]", lines_algo3),
        ("3.3.4. Stylometric Logistic Regression Classifier", 4, "Binary Logistic Regression Classifier", "5D Stylometric Feature Vector v in R^5, Learned Weight Vector w in R^5, Bias intercept beta0", "AI Probability Score S_LR in [0, 100]", lines_algo4),
        ("3.3.5. Character Trigram Cosine Similarity", 5, "Character Trigram Cosine Similarity", "Document text T, Reference profile vectors H_ref, AI_ref in R^20", "AI Probability Score S_Trigram in [0, 100]", lines_algo5),
        ("3.3.6. POS Syntax Ratio Cosine Similarity", 6, "Part-of-Speech Syntax Ratio Classification", "Document text T, Prototype vectors C_AI, C_Human in R^5", "AI Probability Score S_POS in [0, 100]", lines_algo6),
        ("3.3.7. Heuristic Grammar & Typography Analyzer", 7, "Heuristic Grammar & Typography Analysis", "Document text T, Rule set R = {(regex_i, weight_i, name_i)} of 9 weighted patterns", "Grammar Perfection Score G in [0, 100], Detected issue list I", lines_algo7),
        ("3.3.8. Ensemble Scoring & Grammar Calibration", 8, "Weighted Ensemble Scoring & Grammar Calibration", "Component scores {S_LR, S_Trigram, S_KNN, S_POS, S_Cosine, S_Perplexity}, Grammar Perfection Score G", "Final Calibrated AI Probability Score P_final in [0, 100]", lines_algo8)
    ]
    
    # We will iterate through paragraphs to locate target insertion points
    paras = doc.paragraphs
    insert_points = []
    
    for heading_prefix, algo_num, title, inputs, outputs, lines in specs:
        target_idx = None
        for i, p in enumerate(paras):
            if heading_prefix in p.text:
                # Search downstream for the end of text / before next heading
                for j in range(i+1, min(i+15, len(paras))):
                    if paras[j].style.name.startswith('Heading'):
                        target_idx = j - 1
                        break
                if target_idx is None:
                    target_idx = i + 3
                break
        if target_idx is not None:
            insert_points.append((target_idx, algo_num, title, inputs, outputs, lines))
            print(f"Target found for Algo {algo_num} at paragraph {target_idx}: {repr(paras[target_idx].text[:50])}")

    # Sort in reverse index order so insertions don't invalidate remaining indices
    insert_points.sort(key=lambda x: x[0], reverse=True)
    
    for target_idx, algo_num, title, inputs, outputs, lines in insert_points:
        p_target = paras[target_idx]
        tbl = create_ieee_table(doc, algo_num, title, inputs, outputs, lines)
        p_target._element.addnext(tbl._element)
        print(f"Successfully inserted IEEE Table for Algorithm {algo_num}")

    doc.save('Ai text Detector Report_Updated.docx')
    print("\nSAVED UPDATED REPORT DOCUMENT SUCCESSFULLY!")

if __name__ == '__main__':
    main()
