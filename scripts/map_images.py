import docx
import re
import os

doc = docx.Document('ai_text_detector_project_report.docx')
rid_to_target = {}
for rId, rel in doc.part.rels.items():
    if 'image' in rel.target_ref:
        rid_to_target[rId] = rel.target_ref

print("=== Image Mappings ===")
for i, p in enumerate(doc.paragraphs):
    xml = p._element.xml
    if 'embed=' in xml:
        rids = re.findall(r'embed="(rId\d+)"', xml)
        for r in rids:
            target = rid_to_target.get(r, 'Unknown')
            print(f"P{i:03d} -> {r} ({target})")
            for offset in range(-2, 3):
                if 0 <= i+offset < len(doc.paragraphs):
                    t = doc.paragraphs[i+offset].text.strip()
                    if t:
                        print(f"   [{i+offset}]: {t[:90]}")
